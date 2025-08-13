from typing import Any

from django.urls import reverse
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from apps.api import utils

FILE_TYPE: utils.FileTypes = utils.FileTypes.MICROSOFT_WORD


def add_section_divider(document, thickness=1, color="B4B4B4"):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(thickness * 8))  # 1 pt = 8 eighths
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def generate_docx(data: dict[str, Any], meeting_id: str) -> tuple[bool, str | None]:
    with open("/tmp/debug.log", "a") as f:
        f.write("=== DOCX GENERATION START ===\n")

    try:
        export_path = utils.get_export_path()
        filename = utils.generate_filename(meeting_id=meeting_id, file_type=FILE_TYPE)
        file_path = utils.generate_file_path(export_path=export_path, filename=filename)

        with open("/tmp/debug.log", "a") as f:
            f.write(f"DOCX paths generated - file_path: {file_path}\n")

        meeting_metadata = utils.get_meeting_metadata(data=data)
        questions = utils.get_summarized_meeting_question_analysis(data=data)
        key_takeaways = utils.get_summarized_meeting_key_takeaways(data=data)

        if not (meeting_metadata and questions and key_takeaways):
            with open("/tmp/debug.log", "a") as f:
                f.write(
                    f"DOCX FAILED: Validation failed - metadata: {bool(meeting_metadata)}, questions: {bool(questions)}, takeaways: {bool(key_takeaways)}\n"
                )
            return (False, None)

        with open("/tmp/debug.log", "a") as f:
            f.write("DOCX: All validations passed, creating Document object...\n")

        doc = Document()

        with open("/tmp/debug.log", "a") as f:
            f.write("DOCX: Document object created successfully\n")

        # Title page
        title = doc.add_paragraph(meeting_metadata.get("title", ""))
        title.style = "Title"
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        with open("/tmp/debug.log", "a") as f:
            f.write("DOCX: Title added successfully\n")

        desc = doc.add_paragraph(meeting_metadata.get("description", ""))
        desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = desc.runs[0] if desc.runs else desc.add_run()
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(90, 90, 90)
        desc.paragraph_format.space_after = Pt(20)

        with open("/tmp/debug.log", "a") as f:
            f.write("DOCX: Description added successfully\n")

        meta_table = doc.add_table(rows=1, cols=3)
        meta_table.autofit = True
        hdr_cells = meta_table.rows[0].cells

        def set_cell(cell, text):
            cell.text = text
            para = cell.paragraphs[0]
            run = para.runs[0] if para.runs else para.add_run()
            run.font.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_after = Pt(0)
            shading_el = OxmlElement("w:shd")
            shading_el.set(qn("w:fill"), "E6E6E6")
            cell._tc.get_or_add_tcPr().append(shading_el)

        set_cell(hdr_cells[0], f"📅 Date Created: {meeting_metadata.get('date', '')}")
        set_cell(
            hdr_cells[1], f"🕒 Created At: {meeting_metadata.get('time_created', '')}"
        )
        set_cell(hdr_cells[2], f"🎯 Director: {meeting_metadata.get('author', '')}")

        # Remove borders
        tbl_pr = meta_table._tbl.tblPr
        borders = tbl_pr.xpath("./w:tblBorders")
        if borders:
            for border in borders[0]:
                border.set(qn("w:val"), "nil")

        with open("/tmp/debug.log", "a") as f:
            f.write("DOCX: Metadata table added successfully\n")

        doc.add_paragraph()  # spacing

        # Questions
        for i, q in enumerate(questions):
            heading = doc.add_heading(q.get("question", ""), level=2)
            run = heading.runs[0] if heading.runs else heading.add_run()
            run.font.color.rgb = RGBColor(0, 51, 102)  # navy blue

            resp_p = doc.add_paragraph(
                f"Total Responses: {q.get('response_count', '')}"
            )
            run = resp_p.runs[0] if resp_p.runs else resp_p.add_run()
            run.font.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(80, 80, 80)
            resp_p.paragraph_format.space_after = Pt(4)

            summary_p = doc.add_paragraph(q.get("summary", ""))
            summary_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = summary_p.runs[0] if summary_p.runs else summary_p.add_run()
            run.font.size = Pt(11)
            summary_p.paragraph_format.space_after = Pt(12)

            add_section_divider(doc)

            with open("/tmp/debug.log", "a") as f:
                f.write(f"DOCX: Question {i} added successfully\n")

        # Key Takeaways
        kt_heading = doc.add_heading("Key Takeaways", level=2)
        run = kt_heading.runs[0] if kt_heading.runs else kt_heading.add_run()
        run.font.color.rgb = RGBColor(0, 51, 102)
        add_section_divider(doc)

        for i, takeaway in enumerate(key_takeaways, 1):
            p = doc.add_paragraph(f"{i}. {takeaway}", style="List Number")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.runs[0] if p.runs else p.add_run()
            run.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(6)

        with open("/tmp/debug.log", "a") as f:
            f.write("DOCX: Key takeaways section added successfully\n")

        doc.save(str(file_path))

        with open("/tmp/debug.log", "a") as f:
            f.write(f"DOCX: File saved successfully to {file_path}\n")

        download_url = f"{reverse('download-file', kwargs={'filename': filename})}"

        with open("/tmp/debug.log", "a") as f:
            f.write(f"DOCX: Download URL generated: {download_url}\n")
            f.write("=== DOCX GENERATION SUCCESS ===\n")

        return (True, download_url)

    except Exception as e:
        with open("/tmp/debug.log", "a") as f:
            f.write(f"DOCX GENERATION EXCEPTION: {str(e)}\n")
            f.write(f"Exception type: {type(e).__name__}\n")
            f.write("=== DOCX GENERATION FAILED ===\n")
        return (False, None)
